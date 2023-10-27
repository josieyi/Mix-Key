#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os
import re
import abc
import copy
import docx
import math
import pysmiles
import torch.utils.data

import numpy as np
import pandas as pd
import networkx as nx
import scipy.sparse as sp
from rdkit import Chem
from utils.MoleculeConvert import Sanitize_smiles
from utils.MoleculeNetworkCorrect import *
from utils.split import random_scaffold_split, random_split

def data_split(args, dataset):
    split_path = 'DataSplits/{}_{}/'.format(args.bmname, args.split)
    if not os.path.exists(split_path):
        os.makedirs(split_path)
    split_name = split_path + str(args.bmname) + '_' + str(args.seed) + '_split' + '.docx'
    if os.path.exists(split_name) == False:
        smiles_list = pd.read_csv('data/{}.csv'.format(args.bmname), header=None)[0].tolist()
        smiles_list.remove('smiles')
        train_idx = []
        test_idx = []
        valid_idx = []
        if args.split == 'scaffold':
            train_splits, test_splits, val_splits = random_scaffold_split(dataset, smiles_list, task_idx=None,
                                                                          null_value=0, frac_train=0.8, frac_valid=0.1,
                                                                          frac_test=0.1, seed=args.seed)
        elif args.split == 'random':
            train_splits, test_splits, val_splits = random_split(dataset, task_idx=None, null_value=0,
                                                                 frac_train=0.8, frac_valid=0.1, frac_test=0.1, seed=0)

        train_idx.append(train_splits)
        test_idx.append(test_splits)
        valid_idx.append(val_splits)
        file = docx.Document()
        file.add_paragraph(str(args.bmname) + '_' + str(args.seed) + '\n')
        file.add_paragraph('train_splits:' + str(train_idx))
        file.add_paragraph('test_splits:' + str(test_idx))
        file.add_paragraph('val_splits:' + str(valid_idx))
        file.save(split_name)
    else:
        print('read split file')
        split_file = docx.Document(split_name)
        pattern = re.compile(r'\d+')
        strs = split_file.paragraphs[1].text[len('train_splits=['):-1]
        num = 0
        while num < len(strs):
            num1 = strs[num:].index(']') + num
            train_split_str = re.findall(pattern, strs[num:num1 + 1])
            train_split = [int(a) for a in train_split_str]
            train_splits = train_split
            num = num1 + 3
        strs = split_file.paragraphs[2].text[len('test_splits=['):-1]
        num = 0
        while num < len(strs):
            num1 = strs[num:].index(']') + num
            test_split_str = re.findall(pattern, strs[num:num1 + 1])
            test_split = [int(a) for a in test_split_str]
            test_splits = test_split
            num = num1 + 3
        strs = split_file.paragraphs[3].text[len('val_splits=['):-1]
        num = 0
        while num < len(strs):
            num1 = strs[num:].index(']') + num
            val_split_str = re.findall(pattern, strs[num:num1 + 1])
            val_split = [int(a) for a in val_split_str]
            val_splits = val_split
            num = num1 + 3
        print("Split Read Success")

    return train_splits, test_splits, val_splits

def read_data(datadir, dataname, task=0):
    return read_csv_file(datadir, dataname, task)

def read_csv_file(datadir, dataname, task=0):

    file_name = './' + datadir + '/' + dataname + '.csv'
    database = pd.read_csv(file_name)
    column_label = [label for label in list(database.columns)]
    simles_database = list(database[column_label[0]])
    labels_list = [list(database[id]) for id in column_label[1:]]

    assert task < len(column_label[1:]), "Classification task out of sum"
    # smiles:CC1(C)CC(O)CC(C)(C)N1[O];smiles1:CC1(C)CC(O)CC(C)(C)N1O
    if len(column_label[1:]) < 2:
        labels = labels_list[task]
        print("Task: {}".format(task))
    else:
        labels = list(map(list, zip(*labels_list)))
        print("multi_task")

    elements = []
    for id, smiles in enumerate(simles_database):
        mol = Chem.MolFromSmiles(smiles)
        smiles = Chem.MolToSmiles(mol, isomericSmiles=False, allHsExplicit=True)
        if 'se' in smiles:
            smiles = smiles.replace('se', 'Se')
        try:
            graph = pysmiles.read_smiles(smiles, reinterpret_aromatic=False)
        except ValueError:
            print('value error{}'.format(smiles))
        elements.extend(list(nx.get_node_attributes(graph, 'element').values()))
    elements = list(set(elements))
    elements_num = len(elements)
    elements_id = range(elements_num)
    node_indicator = dict(zip(elements_id, elements))

    graphs = []
    for id, smiles in enumerate(simles_database):
        smiles = Sanitize_smiles(smiles)
        mol = Chem.MolFromSmiles(smiles)
        Chem.Kekulize(mol)
        graph = Mol2Graph(mol, node_indicator, H_flag=0)
        mol = Graph2Mol(graph)
        smiles1 = Chem.MolToSmiles(mol, allHsExplicit=True)

        if 'se' in smiles:
            smiles = smiles.replace('se', 'Se')
        smiles1 = Sanitize_smiles(smiles1)

        graph.graph['label'] = labels[id]
        graphs.append(graph)
    if len(column_label[1:]) < 2:
        graphs = [g for g in graphs if (not math.isnan(g.graph['label']))]
    else:
        graphs = [g for g in graphs if len(g.graph['label']) == len(column_label[1:])]
    Graphs = [g for g in graphs if g.number_of_nodes() > 0]
    for i, g in enumerate(Graphs):
        Graphs[i].graph['id'] = i
        Graphs[i].graph['stand'] = 1
        Graphs[i].graph['right'] = 1
    max_num_nodes = max([graph.number_of_nodes() for graph in Graphs])
    return Graphs, max_num_nodes, node_indicator, len(column_label[1:])

def Network2TuDataset(graph, device, edge_feat_dim=4):
    g = copy.deepcopy(graph)
    adj = sp.coo_matrix(nx.adjacency_matrix(g))
    if len(g) == 1:
        edge_index = torch.LongTensor([[0], [0]]).to(device)
        edgetype = np.array([0])
    else:
        edge_index = torch.vstack((torch.LongTensor(adj.row), torch.LongTensor(adj.col))).to(device)
        try:
            edgetype = np.array([g.get_edge_data(edge_index[0][num].item(), edge_index[1][num].item())['order'] for num in range(edge_index.shape[1])])
        except:
            edgetype = []
            for num in range(edge_index.shape[1]):
                try:
                    type = g.get_edge_data(edge_index[0][num].item(), edge_index[1][num].item())['order']
                    edgetype.append(type)
                except:
                    type = 1
                    edgetype.append(type)
            edgetype = np.array(edgetype)
    edge_feat = torch.LongTensor(edgetype).to(device)
    y = torch.tensor([g.graph['label']]).to(device)
    feat = list(nx.get_node_attributes(g, 'label').values())
    x = torch.FloatTensor(feat).to(device)
    return {'edge_index': edge_index, 'edge_feat': edge_feat, 'x': x, 'y': y}